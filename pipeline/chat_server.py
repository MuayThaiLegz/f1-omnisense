"""RAG Chatbot server — F1 Knowledge Base Agent.

Queries MongoDB Atlas vector search for context, then generates
answers via Groq LLM (llama-3.3-70b-versatile).

Usage:
    python pipeline/chat_server.py          # Runs on port 8100
"""

from __future__ import annotations

import os
import sys
import json
import tempfile
from pathlib import Path

import numpy as np

from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# Ensure project root is on path
sys.path.insert(0, str(Path(__file__).parent.parent))
# OmniSuite packages
sys.path.insert(0, str(Path(__file__).parent.parent / "omnisuitef1"))

from dotenv import load_dotenv
load_dotenv(Path(__file__).parent.parent / ".env")

from groq import Groq
from pipeline.vectorstore import AtlasVectorStore
from pipeline.embeddings import NomicEmbedder
from pipeline.model_3d_server import router as model_3d_router, mount_3d_static
from pipeline.omni_health_router import router as omni_health_router
from pipeline.omni_analytics_router import router as omni_analytics_router
from pipeline.omni_rag_router import router as omni_rag_router
from pipeline.opponents.server import router as opponents_router

# ── Config ───────────────────────────────────────────────────────────────

GROQ_MODEL = os.getenv("GROQ_REASONING_MODEL", "llama-3.3-70b-versatile")
PORT = int(os.getenv("PORT", os.getenv("API_PORT", "8100")))
USE_OMNIRAG = os.getenv("USE_OMNIRAG", "").lower() in ("1", "true", "yes")

app = FastAPI(title="F1 OmniSense API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount 3D model generation routes
app.include_router(model_3d_router)
mount_3d_static(app)

# Mount OmniSuite routers
app.include_router(omni_health_router)
app.include_router(omni_analytics_router)
app.include_router(omni_rag_router)
app.include_router(opponents_router)

# Lazy-init singletons
_groq: Groq | None = None
_vs: AtlasVectorStore | None = None
_embedder: NomicEmbedder | None = None
_clip_index: dict | None = None
_clip_embedder = None

CLIP_INDEX_PATH = Path(__file__).parent.parent / "f1data" / "McMedia" / "clip_index.json"


def get_groq() -> Groq:
    global _groq
    if _groq is None:
        _groq = Groq(api_key=os.getenv("GROQ_API_KEY"))
    return _groq


def get_vs() -> AtlasVectorStore:
    global _vs
    if _vs is None:
        _vs = AtlasVectorStore()
    return _vs


def get_embedder() -> NomicEmbedder:
    global _embedder
    if _embedder is None:
        _embedder = NomicEmbedder()
    return _embedder


def get_clip_index() -> dict:
    """Load pre-built CLIP index from disk."""
    global _clip_index
    if _clip_index is None:
        if not CLIP_INDEX_PATH.exists():
            raise FileNotFoundError(
                f"CLIP index not found. Run: python pipeline/clip_index.py"
            )
        with open(CLIP_INDEX_PATH) as f:
            _clip_index = json.load(f)
        # Pre-compute numpy arrays for fast search
        _clip_index["_image_vecs"] = np.array(
            [img["embedding"] for img in _clip_index["images"]]
        )
    return _clip_index


def get_clip_embedder():
    """Lazy-load CLIP embedder for query embedding."""
    global _clip_embedder
    if _clip_embedder is None:
        from pipeline.embeddings import CLIPEmbedder
        _clip_embedder = CLIPEmbedder()
    return _clip_embedder


# ── Models ───────────────────────────────────────────────────────────────

class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str


class ChatRequest(BaseModel):
    message: str
    history: list[ChatMessage] = []


class ChatResponse(BaseModel):
    answer: str
    sources: list[dict]


# ── System Prompt ────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are the F1 OmniSense Knowledge Agent — an expert on Formula 1 technical regulations, car specifications, equipment, and engineering standards.

You have access to a knowledge base extracted from FIA 2024 Technical Regulations and related engineering documents. When answering questions:

1. Use ONLY the provided context to answer. If the context doesn't contain enough information, say so clearly.
2. Cite specific regulation IDs (e.g., "Article 3.5.2") and page numbers when available.
3. Be precise with numerical values, units, and tolerances.
4. For dimensional specifications, always include the value and unit.
5. When discussing equipment, reference tags and types.
6. Keep answers concise but thorough. Use bullet points for lists.
7. If a question is ambiguous, briefly clarify what you're answering.

You are speaking with F1 engineers and technical staff — use appropriate technical language."""


# ── RAG Pipeline ─────────────────────────────────────────────────────────

def _text_search_fallback(query: str, k: int = 8) -> list[dict]:
    """Fallback: keyword search on f1_knowledge when vector search is unavailable."""
    vs = get_vs()
    coll = vs.collection
    keywords = [w for w in query.split() if len(w) > 2]
    if not keywords:
        keywords = query.split()
    regex = "|".join(keywords)
    try:
        results = list(
            coll.find(
                {"page_content": {"$regex": regex, "$options": "i"}},
                {"page_content": 1, "metadata": 1, "_id": 0},
            ).limit(k)
        )
    except Exception:
        results = list(coll.find({}, {"page_content": 1, "metadata": 1, "_id": 0}).limit(k))
    sources = []
    for r in results:
        meta = r.get("metadata", {})
        sources.append({
            "content": r.get("page_content", ""),
            "data_type": meta.get("data_type", ""),
            "category": meta.get("category", ""),
            "source": meta.get("source", ""),
            "page": meta.get("page", 0),
        })
    return sources


def retrieve_context(query: str, k: int = 8) -> list[dict]:
    """Retrieve relevant documents using text search (fast) with optional
    vector search upgrade if embedder is already loaded."""
    global _embedder
    # If embedder is already loaded, use vector search
    if _embedder is not None:
        try:
            vs = get_vs()
            query_vec = _embedder.embed([query])[0]
            docs = vs.similarity_search(query, k=k, query_embedding=query_vec)
            sources = []
            for doc in docs:
                sources.append({
                    "content": doc.page_content,
                    "data_type": doc.metadata.get("data_type", ""),
                    "category": doc.metadata.get("category", ""),
                    "source": doc.metadata.get("source", ""),
                    "page": doc.metadata.get("page", 0),
                })
            return sources
        except Exception as e:
            print(f"  Vector search failed ({e}), falling back to text search")
    # Fast text search — no model loading required
    return _text_search_fallback(query, k)


def build_rag_prompt(query: str, sources: list[dict], history: list[ChatMessage]) -> list[dict]:
    """Build the full prompt with system, context, history, and user query."""
    context_parts = []
    for i, src in enumerate(sources, 1):
        context_parts.append(
            f"[{i}] ({src['data_type']}/{src['category']}) "
            f"Page {src['page']} — {src['source']}\n{src['content']}"
        )
    context_block = "\n\n".join(context_parts)

    messages = [{"role": "system", "content": SYSTEM_PROMPT}]

    # Add conversation history (last 10 messages)
    for msg in history[-10:]:
        messages.append({"role": msg.role, "content": msg.content})

    # User message with context
    user_prompt = f"""CONTEXT FROM KNOWLEDGE BASE:
{context_block}

USER QUESTION:
{query}

Answer based on the context above. Cite regulation IDs and page numbers where applicable."""

    messages.append({"role": "user", "content": user_prompt})
    return messages


# ── Endpoints ────────────────────────────────────────────────────────────

@app.post("/chat", response_model=ChatResponse)
@app.post("/api/chat", response_model=ChatResponse)
def chat(req: ChatRequest):
    """RAG chat endpoint. Delegates to OmniRAG when USE_OMNIRAG=true."""
    if USE_OMNIRAG:
        from pipeline.omni_rag_router import chat as _omni_chat, ChatRequest as _OmniReq
        result = _omni_chat(_OmniReq(message=req.message))
        return ChatResponse(answer=result["answer"], sources=result.get("sources", []))

    # 1. Retrieve context
    sources = retrieve_context(req.message, k=8)

    # 2. Build prompt
    messages = build_rag_prompt(req.message, sources, req.history)

    # 3. Generate answer via Groq
    groq = get_groq()
    completion = groq.chat.completions.create(
        model=GROQ_MODEL,
        messages=messages,
        temperature=0.3,
        max_tokens=2048,
    )

    answer = completion.choices[0].message.content

    return ChatResponse(answer=answer, sources=sources)


@app.get("/health")
def health():
    vs = get_vs()
    return {
        "status": "ok",
        "model": GROQ_MODEL,
        "documents": vs.count(),
        "services": ["knowledge_agent", "3d_model_gen"],
    }


# ── CLIP Visual Search ───────────────────────────────────────────────

@app.get("/visual-search")
def visual_search(q: str, k: int = 8):
    """Search images by text query using CLIP embeddings."""
    index = get_clip_index()
    clip = get_clip_embedder()

    # Embed text query into CLIP space
    query_vec = np.array(clip.embed_text(q))
    query_vec = query_vec / np.linalg.norm(query_vec)

    # Cosine similarity against all image embeddings
    image_vecs = index["_image_vecs"]
    image_norms = image_vecs / np.linalg.norm(image_vecs, axis=1, keepdims=True)
    similarities = query_vec @ image_norms.T

    # Top-k results
    top_indices = np.argsort(similarities)[::-1][:k]

    results = []
    for idx in top_indices:
        img = index["images"][int(idx)]
        results.append({
            "path": img["path"],
            "score": round(float(similarities[idx]), 4),
            "auto_tags": img["auto_tags"],
            "source_video": img["source_video"],
            "frame_index": img["frame_index"],
        })

    return {"query": q, "results": results}


@app.get("/visual-tags")
def visual_tags():
    """Return all images with their auto-tags for gallery filtering."""
    index = get_clip_index()

    images = []
    for img in index["images"]:
        images.append({
            "path": img["path"],
            "auto_tags": img["auto_tags"],
            "source_video": img["source_video"],
            "frame_index": img["frame_index"],
        })

    # Collect all unique tag labels with their max scores
    tag_summary: dict[str, float] = {}
    for img in index["images"]:
        for tag in img["auto_tags"]:
            label = tag["label"]
            if label not in tag_summary or tag["score"] > tag_summary[label]:
                tag_summary[label] = tag["score"]

    top_tags = sorted(tag_summary.items(), key=lambda x: -x[1])

    return {
        "images": images,
        "tags": [{"label": t[0], "max_score": round(t[1], 4)} for t in top_tags],
        "stats": index["stats"],
    }


# ── Document Upload & Ingestion ──────────────────────────────────────────

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".json", ".md"}


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    import fitz
    doc = fitz.open(file_path)
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_from_plain(file_path: str) -> str:
    """Read plain text files (txt, csv, json, md)."""
    return Path(file_path).read_text(encoding="utf-8", errors="replace")


EXTRACTORS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".txt": extract_text_from_plain,
    ".csv": extract_text_from_plain,
    ".json": extract_text_from_plain,
    ".md": extract_text_from_plain,
}


@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload a document, extract text, embed, and ingest into the knowledge base."""
    filename = file.filename or "unknown"
    ext = Path(filename).suffix.lower()

    if ext not in ALLOWED_EXTENSIONS:
        return {"filename": filename, "status": "error",
                "error": f"Unsupported format: {ext}. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"}

    # Save to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        # Extract text
        extractor = EXTRACTORS[ext]
        text = extractor(tmp_path)

        if not text.strip():
            return {"filename": filename, "status": "error", "error": "No text extracted from file"}

        # Chunk text
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = splitter.split_text(text)

        if not chunks:
            return {"filename": filename, "status": "error", "error": "No chunks created"}

        # Build LangChain Documents
        from langchain_core.documents import Document
        docs = []
        for i, chunk in enumerate(chunks):
            docs.append(Document(
                page_content=chunk,
                metadata={
                    "data_type": "uploaded_document",
                    "category": "user_upload",
                    "source": filename,
                    "chunk": i + 1,
                    "total_chunks": len(chunks),
                },
            ))

        # Embed
        embedder = get_embedder()
        texts = [doc.page_content for doc in docs]
        embeddings: list[list[float]] = []
        for i in range(0, len(texts), 32):
            batch = texts[i:i + 32]
            embeddings.extend(embedder.embed(batch))

        # Upsert to Atlas
        vs = get_vs()
        count = vs.upsert_documents(docs, embeddings)

        return {
            "filename": filename,
            "status": "ok",
            "chunks": count,
            "text_length": len(text),
        }

    except Exception as e:
        return {"filename": filename, "status": "error", "error": str(e)}
    finally:
        os.unlink(tmp_path)


# ── Data API — serve MongoDB collections to frontend ─────────────────────

from pymongo import MongoClient as _MongoClient

_data_client: _MongoClient | None = None
_data_db = None

def get_data_db():
    global _data_client, _data_db
    if _data_db is None:
        _data_client = _MongoClient(os.getenv("MONGODB_URI", ""))
        _data_db = _data_client[os.getenv("MONGODB_DB", "McLaren_f1")]
    return _data_db

@app.get("/api/local/jolpica/race_results")
async def jolpica_race_results():
    db = get_data_db()
    docs = list(db["race_results"].find({}, {"_id": 0}))
    return docs

@app.get("/api/local/jolpica/driver_standings")
async def jolpica_driver_standings():
    db = get_data_db()
    races = list(db["race_results"].find({}, {"_id": 0}).sort([("season", 1), ("round", 1)]))
    # Build standings by aggregating points per driver per season
    from collections import defaultdict
    driver_data = defaultdict(lambda: {"points": 0, "wins": 0, "Driver": None, "Constructors": [], "season": ""})
    for race in races:
        season = race.get("season", "")
        for result in race.get("Results", []):
            driver = result.get("Driver", {})
            did = driver.get("driverId", "")
            key = f"{season}_{did}"
            d = driver_data[key]
            d["points"] += float(result.get("points", 0))
            if result.get("position") == "1":
                d["wins"] += 1
            d["Driver"] = driver
            d["season"] = season
            constructor = result.get("Constructor", {})
            if constructor and constructor not in d["Constructors"]:
                d["Constructors"] = [constructor]
    # Format as JolpicaDriverStanding[]
    standings = []
    for key, d in driver_data.items():
        if d["Driver"]:
            standings.append({
                "position": "0",
                "positionText": "0",
                "points": str(d["points"]),
                "wins": str(d["wins"]),
                "Driver": d["Driver"],
                "Constructors": d["Constructors"],
                "season": d["season"],
            })
    # Sort by season desc, points desc
    standings.sort(key=lambda x: (-int(x["season"] or "0"), -float(x["points"])))
    # Assign positions per season
    current_season = None
    pos = 0
    for s in standings:
        if s["season"] != current_season:
            current_season = s["season"]
            pos = 1
        s["position"] = str(pos)
        s["positionText"] = str(pos)
        pos += 1
    return standings

@app.get("/api/local/jolpica/constructor_standings")
async def jolpica_constructor_standings():
    db = get_data_db()
    races = list(db["race_results"].find({}, {"_id": 0}).sort([("season", 1), ("round", 1)]))
    from collections import defaultdict
    constructor_data = defaultdict(lambda: {"points": 0, "wins": 0, "Constructor": None, "season": ""})
    for race in races:
        season = race.get("season", "")
        for result in race.get("Results", []):
            constructor = result.get("Constructor", {})
            cid = constructor.get("constructorId", "")
            key = f"{season}_{cid}"
            d = constructor_data[key]
            d["points"] += float(result.get("points", 0))
            if result.get("position") == "1":
                d["wins"] += 1
            d["Constructor"] = constructor
            d["season"] = season
    standings = []
    for key, d in constructor_data.items():
        if d["Constructor"]:
            standings.append({
                "position": "0",
                "positionText": "0",
                "points": str(d["points"]),
                "wins": str(d["wins"]),
                "Constructor": d["Constructor"],
                "season": d["season"],
            })
    standings.sort(key=lambda x: (-int(x["season"] or "0"), -float(x["points"])))
    current_season = None
    pos = 0
    for s in standings:
        if s["season"] != current_season:
            current_season = s["season"]
            pos = 1
        s["position"] = str(pos)
        s["positionText"] = str(pos)
        pos += 1
    return standings

@app.get("/api/local/jolpica/qualifying")
async def jolpica_qualifying():
    return []

@app.get("/api/local/jolpica/circuits")
async def jolpica_circuits():
    return []

@app.get("/api/local/jolpica/pit_stops")
async def jolpica_pit_stops():
    return []

@app.get("/api/local/jolpica/lap_times")
async def jolpica_lap_times():
    return []

@app.get("/api/local/jolpica/drivers")
async def jolpica_drivers():
    return []

@app.get("/api/local/jolpica/seasons")
async def jolpica_seasons():
    return []

@app.get("/api/local/pipeline/anomaly")
async def pipeline_anomaly():
    db = get_data_db()
    snapshot = db["anomaly_scores_snapshot"].find_one({}, {"_id": 0})
    return snapshot or {}

@app.get("/api/local/pipeline/intelligence")
async def pipeline_intelligence():
    """Serve regulations, equipment, dimensions, materials from f1_knowledge."""
    db = get_data_db()
    rules = []
    equipment = []
    dimensional_data = []
    material_specs = []

    for doc in db["f1_knowledge"].find({}, {"_id": 0, "embedding": 0}):
        meta = doc.get("metadata", {})
        dt = meta.get("data_type", "")
        content = doc.get("page_content", "")

        if dt == "regulation":
            # Parse description from page_content
            desc_lines = content.split("\n")
            description = desc_lines[1] if len(desc_lines) > 1 else content
            rules.append({
                "id": meta.get("rule_id", ""),
                "category": meta.get("category", ""),
                "description": description,
                "value": None,
                "unit": None,
                "condition": None,
                "reference": None,
                "severity": meta.get("severity", "info"),
                "source_standard": meta.get("source", ""),
                "_source": meta.get("source", ""),
                "_page": meta.get("page", 0),
            })
        elif dt == "equipment":
            # Parse equipment fields from page_content
            lines = content.split("\n")
            eq_type = ""
            eq_desc = ""
            eq_location = None
            for line in lines:
                if line.startswith("Type: "):
                    eq_type = line[6:]
                elif line.startswith("Location: "):
                    eq_location = line[10:]
                elif not line.startswith("["):
                    eq_desc = line
            equipment.append({
                "tag": meta.get("tag", ""),
                "type": eq_type or meta.get("category", ""),
                "description": eq_desc,
                "kks": "",
                "specs": {},
                "location_description": eq_location,
                "_source": meta.get("source", ""),
                "_page": meta.get("page", 0),
            })
        elif dt == "dimension":
            # Parse dimensional data from page_content
            lines = content.split("\n")
            dimension_desc = lines[1] if len(lines) > 1 else ""
            value = None
            unit = ""
            for line in lines:
                if line.startswith("Value: "):
                    val_str = line[7:].strip()
                    parts = val_str.split()
                    try:
                        value = float(parts[0])
                        unit = " ".join(parts[1:]) if len(parts) > 1 else ""
                    except (ValueError, IndexError):
                        value = val_str
            dimensional_data.append({
                "component": meta.get("component", ""),
                "dimension": dimension_desc,
                "value": value,
                "unit": unit,
                "_source": meta.get("source", ""),
                "_page": meta.get("page", 0),
            })
        elif dt == "material":
            lines = content.split("\n")
            application = ""
            for line in lines:
                if line.startswith("Application: "):
                    application = line[13:]
            material_specs.append({
                "material": meta.get("material", ""),
                "application": application,
                "properties": {},
                "_source": meta.get("source", ""),
                "_page": meta.get("page", 0),
            })

    return {
        "documents": [],
        "rules": rules,
        "equipment": equipment,
        "dimensional_data": dimensional_data,
        "material_specs": material_specs,
        "stats": {
            "total_pages": 0,
            "total_rules": len(rules),
            "total_equipment": len(equipment),
            "total_dimensions": len(dimensional_data),
            "total_materials": len(material_specs),
            "total_tokens_in": 0,
            "total_tokens_out": 0,
            "total_cost_usd": 0,
            "total_latency_s": 0,
        },
    }

@app.get("/api/local/pipeline/gdino")
async def pipeline_gdino():
    return {}

@app.get("/api/local/pipeline/fused")
async def pipeline_fused():
    return {}

@app.get("/api/local/pipeline/minicpm")
async def pipeline_minicpm():
    return {}

@app.get("/api/local/pipeline/videomae")
async def pipeline_videomae():
    return {}

@app.get("/api/local/pipeline/timesformer")
async def pipeline_timesformer():
    return {}

def _build_session_map():
    """Build mapping from _source_file → session_key and (year, race) → session_key.

    Returns (src_to_key, year_race_to_key) where year_race_to_key keys are
    ``"YYYY|Race Name"`` strings so that 2023 and 2024 races get distinct
    session keys.
    """
    db = get_data_db()
    sources = db["telemetry"].distinct("_source_file")
    src_to_key = {}
    year_race_to_key = {}  # "2024|Monaco Grand Prix" → session_key
    session_key = 9000
    for src in sorted(sources):
        parts = src.replace(".csv", "").split("_")
        if len(parts) < 3:
            continue
        year = parts[0]
        race_name = " ".join(parts[1:]).replace(" Race", "")
        src_to_key[src] = session_key
        year_race_to_key[f"{year}|{race_name}"] = session_key
        session_key += 1
    return src_to_key, year_race_to_key


def _resolve_sk(year_race_to_key: dict, year: str, race: str) -> int:
    """Look up session_key for a (year, race) pair with fallbacks."""
    key = f"{year}|{race}"
    if key in year_race_to_key:
        return year_race_to_key[key]
    # Try adding "Grand Prix" suffix
    key2 = f"{year}|{race} Grand Prix"
    if key2 in year_race_to_key:
        return year_race_to_key[key2]
    return 9000

_DRIVER_NUMBERS = {"NOR": 4, "PIA": 81}
_DRIVER_META = {
    "NOR": {"number": 4, "full_name": "Lando NORRIS", "broadcast_name": "L NORRIS",
            "first_name": "Lando", "last_name": "Norris", "team_name": "McLaren",
            "team_colour": "FF8000", "country_code": "GBR", "headshot_url": ""},
    "PIA": {"number": 81, "full_name": "Oscar PIASTRI", "broadcast_name": "O PIASTRI",
            "first_name": "Oscar", "last_name": "Piastri", "team_name": "McLaren",
            "team_colour": "FF8000", "country_code": "AUS", "headshot_url": ""},
}

# Map CSV GP names to OpenF1 circuit_short_name values
_GP_TO_CIRCUIT: dict[str, str] = {
    "Abu Dhabi Grand Prix": "Yas Marina Circuit",
    "Australian Grand Prix": "Melbourne",
    "Azerbaijan Grand Prix": "Baku",
    "Bahrain Grand Prix": "Sakhir",
    "Belgian Grand Prix": "Spa-Francorchamps",
    "British Grand Prix": "Silverstone",
    "Canadian Grand Prix": "Montreal",
    "Chinese Grand Prix": "Shanghai",
    "Dutch Grand Prix": "Zandvoort",
    "Emilia Romagna Grand Prix": "Imola",
    "Hungarian Grand Prix": "Hungaroring",
    "Italian Grand Prix": "Monza",
    "Japanese Grand Prix": "Suzuka",
    "Las Vegas Grand Prix": "Las Vegas",
    "Mexico City Grand Prix": "Mexico City",
    "Miami Grand Prix": "Miami",
    "Monaco Grand Prix": "Monte Carlo",
    "Qatar Grand Prix": "Lusail",
    "Saudi Arabian Grand Prix": "Jeddah",
    "Singapore Grand Prix": "Singapore",
    "Spanish Grand Prix": "Catalunya",
    "Austrian Grand Prix": "Spielberg",
    "United States Grand Prix": "Austin",
    "São Paulo Grand Prix": "Interlagos",
    "Brazilian Grand Prix": "Interlagos",
}

_GP_TO_COUNTRY: dict[str, str] = {
    "Abu Dhabi Grand Prix": "UAE",
    "Australian Grand Prix": "Australia",
    "Azerbaijan Grand Prix": "Azerbaijan",
    "Bahrain Grand Prix": "Bahrain",
    "Belgian Grand Prix": "Belgium",
    "British Grand Prix": "Great Britain",
    "Canadian Grand Prix": "Canada",
    "Chinese Grand Prix": "China",
    "Dutch Grand Prix": "Netherlands",
    "Emilia Romagna Grand Prix": "Italy",
    "Hungarian Grand Prix": "Hungary",
    "Italian Grand Prix": "Italy",
    "Japanese Grand Prix": "Japan",
    "Las Vegas Grand Prix": "United States",
    "Mexico City Grand Prix": "Mexico",
    "Miami Grand Prix": "United States",
    "Monaco Grand Prix": "Monaco",
    "Qatar Grand Prix": "Qatar",
    "Saudi Arabian Grand Prix": "Saudi Arabia",
    "Singapore Grand Prix": "Singapore",
    "Spanish Grand Prix": "Spain",
    "Austrian Grand Prix": "Austria",
    "United States Grand Prix": "United States",
    "São Paulo Grand Prix": "Brazil",
    "Brazilian Grand Prix": "Brazil",
}

@app.get("/api/local/openf1/sessions")
async def openf1_sessions():
    """Generate OpenF1-format sessions from telemetry data."""
    db = get_data_db()
    sources = db["telemetry"].distinct("_source_file")
    sessions = []
    session_key = 9000
    race_idx = 0
    prev_year = None
    for src in sorted(sources):
        parts = src.replace(".csv", "").split("_")
        if len(parts) < 3:
            continue
        year = parts[0]
        if year != prev_year:
            race_idx = 0
            prev_year = year
        race_name = " ".join(parts[1:]).replace(" Race", "")
        circuit_short = _GP_TO_CIRCUIT.get(race_name, race_name.split()[0])
        country = _GP_TO_COUNTRY.get(race_name, circuit_short)
        month = 3 + race_idx
        if month > 12:
            month = 12
        sessions.append({
            "session_key": session_key,
            "session_name": "Race",
            "session_type": "Race",
            "date_start": f"{year}-{month:02d}-15T14:00:00",
            "date_end": f"{year}-{month:02d}-15T16:00:00",
            "year": int(year),
            "circuit_key": session_key,
            "circuit_short_name": circuit_short,
            "country_name": country,
            "country_key": session_key,
            "location": circuit_short,
            "meeting_key": session_key,
            "meeting_name": f"{year} {race_name}",
            "_source_file": src,
        })
        session_key += 1
        race_idx += 1
    return sessions

@app.get("/api/local/openf1/drivers")
async def openf1_drivers():
    """Generate OpenF1-format driver data for all sessions."""
    _, yr_to_key = _build_session_map()
    drivers = []
    for sk in yr_to_key.values():
        for code, meta in _DRIVER_META.items():
            drivers.append({
                "session_key": sk, "meeting_key": sk,
                "driver_number": meta["number"],
                "broadcast_name": meta["broadcast_name"],
                "full_name": meta["full_name"],
                "name_acronym": code,
                "team_name": meta["team_name"],
                "team_colour": meta["team_colour"],
                "first_name": meta["first_name"],
                "last_name": meta["last_name"],
                "country_code": meta["country_code"],
                "headshot_url": meta["headshot_url"],
            })
    return drivers

@app.get("/api/local/openf1/laps")
async def openf1_laps():
    """Generate OpenF1-format lap data from telemetry."""
    db = get_data_db()
    _, yr_to_key = _build_session_map()
    pipeline_agg = [
        {"$group": {
            "_id": {"Driver": "$Driver", "Year": "$Year", "Race": "$Race", "LapNumber": "$LapNumber"},
            "lap_time": {"$first": "$LapTime"},
            "top_speed": {"$max": "$Speed"},
            "date": {"$first": "$Date"},
        }},
        {"$sort": {"_id.Year": 1, "_id.Race": 1, "_id.LapNumber": 1}},
    ]
    results = list(db["telemetry"].aggregate(pipeline_agg, allowDiskUse=True))
    laps = []
    for r in results:
        ident = r["_id"]
        driver = ident.get("Driver", "")
        year = ident.get("Year", "")
        race = ident.get("Race", "")
        lap_num = ident.get("LapNumber")
        if lap_num is None:
            continue
        sk = _resolve_sk(yr_to_key, year, race)
        lap_duration = None
        lt = r.get("lap_time", "")
        if lt and "days" in str(lt):
            try:
                time_part = str(lt).split(" ")[-1]
                parts = time_part.split(":")
                lap_duration = float(parts[0]) * 3600 + float(parts[1]) * 60 + float(parts[2])
            except (ValueError, IndexError):
                pass
        laps.append({
            "session_key": sk, "meeting_key": sk,
            "driver_number": _DRIVER_NUMBERS.get(driver, 0),
            "lap_number": int(lap_num),
            "lap_duration": lap_duration,
            "duration_sector_1": None, "duration_sector_2": None, "duration_sector_3": None,
            "is_pit_out_lap": False,
            "date_start": str(r.get("date", "")),
            "st_speed": float(r.get("top_speed", 0)) if r.get("top_speed") else None,
        })
    return laps

@app.get("/api/local/openf1/position")
async def openf1_positions():
    """Generate OpenF1-format position data from telemetry."""
    db = get_data_db()
    _, yr_to_key = _build_session_map()
    pipeline_agg = [
        {"$group": {
            "_id": {"Driver": "$Driver", "Year": "$Year", "Race": "$Race", "LapNumber": "$LapNumber"},
            "date": {"$last": "$Date"},
        }},
        {"$sort": {"_id.Year": 1, "_id.Race": 1, "_id.LapNumber": 1}},
    ]
    results = list(db["telemetry"].aggregate(pipeline_agg, allowDiskUse=True))
    from collections import defaultdict
    lap_groups = defaultdict(list)
    for r in results:
        ident = r["_id"]
        year = ident.get("Year", "")
        race = ident.get("Race", "")
        key = f'{year}_{race}_{ident.get("LapNumber", 0)}'
        lap_groups[key].append({
            "driver": ident.get("Driver", ""),
            "date": r.get("date", ""),
            "year": year,
            "race": race,
        })
    positions = []
    for key, drivers_in_lap in lap_groups.items():
        for pos_idx, d in enumerate(drivers_in_lap):
            sk = _resolve_sk(yr_to_key, d["year"], d["race"])
            positions.append({
                "session_key": sk, "meeting_key": sk,
                "driver_number": _DRIVER_NUMBERS.get(d["driver"], 0),
                "date": str(d["date"]),
                "position": pos_idx + 1,
            })
    return positions

@app.get("/api/local/openf1/weather")
async def openf1_weather():
    """Return weather data for all sessions."""
    _, yr_to_key = _build_session_map()
    weather = []
    for yr_race, sk in yr_to_key.items():
        year = yr_race.split("|")[0]
        weather.append({
            "session_key": sk, "meeting_key": sk,
            "date": f"{year}-06-01T14:00:00", "air_temperature": 25.0,
            "track_temperature": 40.0, "humidity": 55, "pressure": 1013.0,
            "rainfall": False, "wind_direction": 180, "wind_speed": 3.5,
        })
    return weather

@app.get("/api/local/openf1/intervals")
async def openf1_intervals():
    return []

@app.get("/api/local/openf1/pit")
async def openf1_pit():
    """Generate pit stop data from telemetry stint boundaries."""
    db = get_data_db()
    _, yr_to_key = _build_session_map()
    pipeline_agg = [
        {"$match": {"Compound": {"$ne": None}}},
        {"$group": {
            "_id": {"Driver": "$Driver", "Year": "$Year", "Race": "$Race", "Compound": "$Compound"},
            "min_lap": {"$min": "$LapNumber"},
        }},
        {"$sort": {"_id.Race": 1, "_id.Driver": 1, "min_lap": 1}},
    ]
    results = list(db["telemetry"].aggregate(pipeline_agg))
    from collections import defaultdict
    driver_race_stints = defaultdict(list)
    for r in results:
        ident = r["_id"]
        key = f'{ident["Year"]}_{ident["Driver"]}_{ident["Race"]}'
        driver_race_stints[key].append({
            "driver": ident["Driver"],
            "year": ident.get("Year", ""),
            "race": ident["Race"],
            "min_lap": r.get("min_lap", 0) or 0,
        })
    pits = []
    for key, stints in driver_race_stints.items():
        stints.sort(key=lambda s: s["min_lap"])
        for i in range(1, len(stints)):
            driver = stints[i]["driver"]
            year = stints[i]["year"]
            race = stints[i]["race"]
            sk = _resolve_sk(yr_to_key, year, race)
            pits.append({
                "session_key": sk, "meeting_key": sk,
                "driver_number": _DRIVER_NUMBERS.get(driver, 0),
                "date": "", "lap_number": int(stints[i]["min_lap"]),
                "pit_duration": 23.5,
            })
    return pits

@app.get("/api/local/openf1/stints")
async def openf1_stints():
    """Generate stint data from telemetry."""
    db = get_data_db()
    _, yr_to_key = _build_session_map()
    pipeline_agg = [
        {"$match": {"Compound": {"$ne": None}}},
        {"$group": {
            "_id": {"Driver": "$Driver", "Year": "$Year", "Race": "$Race", "Compound": "$Compound"},
            "start_lap": {"$min": "$LapNumber"},
            "end_lap": {"$max": "$LapNumber"},
            "tyre_life": {"$min": "$TyreLife"},
        }},
        {"$sort": {"_id.Year": 1, "_id.Race": 1, "_id.Driver": 1, "start_lap": 1}},
    ]
    results = list(db["telemetry"].aggregate(pipeline_agg, allowDiskUse=True))
    stints = []
    from collections import defaultdict
    stint_counter = defaultdict(int)
    for r in results:
        ident = r["_id"]
        driver = ident.get("Driver", "")
        year = ident.get("Year", "")
        race = ident.get("Race", "")
        key = f"{year}_{driver}_{race}"
        stint_counter[key] += 1
        sk = _resolve_sk(yr_to_key, year, race)
        stints.append({
            "session_key": sk, "meeting_key": sk,
            "driver_number": _DRIVER_NUMBERS.get(driver, 0),
            "stint_number": stint_counter[key],
            "lap_start": int(r.get("start_lap", 0)) if r.get("start_lap") else 0,
            "lap_end": int(r.get("end_lap", 0)) if r.get("end_lap") else 0,
            "compound": ident.get("Compound", "UNKNOWN"),
            "tyre_age_at_start": int(r.get("tyre_life", 0)) if r.get("tyre_life") else 0,
        })
    return stints

@app.get("/api/local/openf1/race_control")
async def openf1_race_control():
    """Return race control events for all sessions."""
    _, yr_to_key = _build_session_map()
    events = []
    for yr_race, sk in yr_to_key.items():
        year = yr_race.split("|")[0]
        events.append({
            "date": f"{year}-06-01T14:00:00", "session_key": sk, "meeting_key": sk,
            "driver_number": None, "lap_number": 1, "category": "Flag",
            "flag": "GREEN", "scope": "Track", "message": "GREEN LIGHT - PIT EXIT OPEN",
        })
    return events

@app.get("/api/local/openf1/{collection}")
async def openf1_other(collection: str):
    """Catch-all for other OpenF1 collections."""
    return []

def _aggregate_telemetry_summary(docs: list[dict]) -> list[dict]:
    """Aggregate raw telemetry docs into CarSummary format grouped by race."""
    from collections import defaultdict
    races = defaultdict(lambda: {
        "speeds": [], "rpms": [], "throttles": [],
        "brake_count": 0, "drs_count": 0, "total": 0,
        "compounds": set(),
    })
    for d in docs:
        race = d.get("Race", "Unknown")
        r = races[race]
        speed = d.get("Speed")
        if speed is not None:
            try:
                r["speeds"].append(float(speed))
            except (ValueError, TypeError):
                pass
        rpm = d.get("RPM")
        if rpm is not None:
            try:
                r["rpms"].append(float(rpm))
            except (ValueError, TypeError):
                pass
        throttle = d.get("Throttle")
        if throttle is not None:
            try:
                r["throttles"].append(float(throttle))
            except (ValueError, TypeError):
                pass
        brake = d.get("Brake")
        if brake is True or brake == "True" or brake == 1:
            r["brake_count"] += 1
        drs = d.get("DRS")
        try:
            if drs is not None and int(float(str(drs))) >= 10:
                r["drs_count"] += 1
        except (ValueError, TypeError):
            pass
        compound = d.get("Compound")
        if compound:
            r["compounds"].add(str(compound))
        r["total"] += 1

    summaries = []
    for race_name, r in races.items():
        if not r["speeds"]:
            continue
        n = r["total"] or 1
        # Extract short race name (remove "Grand Prix" suffix for display)
        short_name = race_name.replace(" Grand Prix", "")
        summaries.append({
            "race": short_name,
            "avgSpeed": round(sum(r["speeds"]) / len(r["speeds"]), 2) if r["speeds"] else 0,
            "topSpeed": round(max(r["speeds"]), 1) if r["speeds"] else 0,
            "avgRPM": round(sum(r["rpms"]) / len(r["rpms"])) if r["rpms"] else 0,
            "maxRPM": round(max(r["rpms"])) if r["rpms"] else 0,
            "avgThrottle": round(sum(r["throttles"]) / len(r["throttles"]), 2) if r["throttles"] else 0,
            "brakePct": round(r["brake_count"] / n * 100, 2),
            "drsPct": round(r["drs_count"] / n * 100, 2),
            "compounds": sorted(r["compounds"]),
            "samples": n,
        })
    summaries.sort(key=lambda x: x["race"])
    return summaries

@app.get("/api/local/mccar-summary/{year}/{driver}")
async def mccar_summary(year: str, driver: str):
    """Aggregate telemetry into CarSummary[] format grouped by race."""
    db = get_data_db()
    docs = list(db["telemetry"].find(
        {"Driver": driver, "Year": year},
        {"_id": 0, "Speed": 1, "RPM": 1, "Throttle": 1, "Brake": 1,
         "DRS": 1, "Compound": 1, "Race": 1}
    ))
    return _aggregate_telemetry_summary(docs)

@app.get("/api/local/mcdriver-summary/{year}/{driver}")
async def mcdriver_summary(year: str, driver: str):
    """Aggregate telemetry into RaceSummary[] format for driver biometrics."""
    db = get_data_db()
    docs = list(db["telemetry"].find(
        {"Driver": driver, "Year": year},
        {"_id": 0, "Speed": 1, "RPM": 1, "Throttle": 1, "Brake": 1,
         "DRS": 1, "Race": 1}
    ))
    from collections import defaultdict
    races = defaultdict(lambda: {"speeds": [], "rpms": [], "total": 0})
    for d in docs:
        race = d.get("Race", "Unknown")
        r = races[race]
        speed = d.get("Speed")
        if speed is not None:
            try:
                r["speeds"].append(float(speed))
            except (ValueError, TypeError):
                pass
        rpm = d.get("RPM")
        if rpm is not None:
            try:
                r["rpms"].append(float(rpm))
            except (ValueError, TypeError):
                pass
        r["total"] += 1
    summaries = []
    for race_name, r in races.items():
        if not r["speeds"]:
            continue
        short_name = race_name.replace(" Grand Prix", "")
        avg_speed = sum(r["speeds"]) / len(r["speeds"]) if r["speeds"] else 0
        # Simulate biometric data from telemetry intensity
        battle_intensity = min(100, round(avg_speed / 3.5, 1))
        summaries.append({
            "race": short_name,
            "avgHR": round(140 + battle_intensity * 0.3, 1),
            "peakHR": round(160 + battle_intensity * 0.4, 1),
            "avgTemp": 36.8,
            "battleIntensity": battle_intensity,
            "airTemp": 25.0,
            "trackTemp": 40.0,
            "samples": r["total"],
        })
    summaries.sort(key=lambda x: x["race"])
    return summaries

def _telemetry_to_csv(docs: list[dict]) -> str:
    """Convert telemetry documents to CSV string."""
    if not docs:
        return ""
    headers = ["Date", "RPM", "Speed", "nGear", "Throttle", "Brake", "DRS",
               "Source", "Time", "SessionTime", "Distance", "Driver", "Year",
               "Race", "LapNumber", "LapTime", "Compound", "TyreLife"]
    lines = [",".join(headers)]
    for d in docs:
        row = []
        for h in headers:
            val = d.get(h, "")
            # Escape commas in values
            s = str(val) if val is not None else ""
            if "," in s:
                s = f'"{s}"'
            row.append(s)
        lines.append(",".join(row))
    return "\n".join(lines)


def _biometrics_to_csv(docs: list[dict]) -> str:
    """Convert biometrics documents to CSV string."""
    if not docs:
        return ""
    headers = ["Date", "RPM", "Speed", "nGear", "Throttle", "Brake", "DRS",
               "Source", "Time", "SessionTime", "Distance", "Driver", "Year",
               "Race", "LapNumber", "LapTime", "Compound", "TyreLife",
               "HeartRate_bpm", "CockpitTemp_C", "AirTemp_C", "TrackTemp_C",
               "Humidity_pct", "BattleIntensity"]
    lines = [",".join(headers)]
    for d in docs:
        row = []
        for h in headers:
            val = d.get(h, "")
            s = str(val) if val is not None else ""
            if "," in s:
                s = f'"{s}"'
            row.append(s)
        lines.append(",".join(row))
    return "\n".join(lines)


@app.get("/api/local/mccar/{year}/{filename}")
async def mccar_csv(year: str, filename: str):
    """Serve telemetry as CSV for a specific race."""
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    # filename like 2024_Abu_Dhabi_Grand_Prix_Race.csv
    source_file = filename.replace(".csv", "") + ".csv"
    docs = list(db["telemetry"].find(
        {"_source_file": source_file},
        {"_id": 0}
    ))
    return PlainTextResponse(_telemetry_to_csv(docs))

@app.get("/api/local/mcdriver/{year}/{filename}")
async def mcdriver_csv(year: str, filename: str):
    """Serve driver biometrics as CSV from the biometrics collection."""
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    # Normalize filename to match _source_file in biometrics collection
    base = filename.replace(".csv", "")
    if not base.endswith("_biometrics"):
        base += "_biometrics"
    source_file = base + ".csv"
    docs = list(db["biometrics"].find(
        {"_source_file": source_file},
        {"_id": 0}
    ))
    if not docs:
        # Fallback: try telemetry collection without _biometrics suffix
        fallback_file = filename.replace(".csv", "").replace("_biometrics", "") + ".csv"
        docs = list(db["telemetry"].find(
            {"_source_file": fallback_file},
            {"_id": 0}
        ))
        return PlainTextResponse(_telemetry_to_csv(docs))
    return PlainTextResponse(_biometrics_to_csv(docs))

@app.get("/api/local/mcracecontext/{year}/tire_stints.csv")
async def mcracecontext_tire_stints(year: str):
    """Generate tire stints CSV from telemetry data."""
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    # Aggregate stint data from telemetry
    pipeline_agg = [
        {"$match": {"Year": year}},
        {"$group": {
            "_id": {"Driver": "$Driver", "Race": "$Race", "Compound": "$Compound"},
            "start_lap": {"$min": "$LapNumber"},
            "end_lap": {"$max": "$LapNumber"},
            "tyre_life": {"$max": "$TyreLife"},
        }},
        {"$sort": {"_id.Race": 1, "_id.Driver": 1, "start_lap": 1}},
    ]
    results = list(db["telemetry"].aggregate(pipeline_agg))
    headers = ["Driver", "Race", "Compound", "StartLap", "EndLap", "TyreLife"]
    lines = [",".join(headers)]
    for r in results:
        ident = r["_id"]
        lines.append(",".join([
            str(ident.get("Driver", "")),
            str(ident.get("Race", "")),
            str(ident.get("Compound", "")),
            str(int(r.get("start_lap", 0))) if r.get("start_lap") else "0",
            str(int(r.get("end_lap", 0))) if r.get("end_lap") else "0",
            str(int(r.get("tyre_life", 0))) if r.get("tyre_life") else "0",
        ]))
    return PlainTextResponse("\n".join(lines))


@app.get("/api/models/{filename}")
async def serve_glb_model(filename: str):
    """Serve GLB 3D model files from MongoDB GridFS with streaming."""
    import gridfs
    from starlette.responses import StreamingResponse, Response
    db = get_data_db()
    fs = gridfs.GridFS(db)
    try:
        grid_file = fs.find_one({"filename": filename})
        if grid_file is None:
            return Response(content="Model not found", status_code=404)

        def stream_chunks():
            while True:
                chunk = grid_file.read(256 * 1024)  # 256KB chunks
                if not chunk:
                    break
                yield chunk

        return StreamingResponse(
            stream_chunks(),
            media_type="model/gltf-binary",
            headers={
                "Cache-Control": "public, max-age=604800",
                "Content-Disposition": f'inline; filename="{filename}"',
                "Content-Length": str(grid_file.length),
            },
        )
    except Exception as e:
        return Response(content=str(e), status_code=500)


@app.get("/api/local/mccsv/driver_career")
async def mccsv_driver_career():
    """Generate driver career CSV from race_results."""
    from starlette.responses import PlainTextResponse
    from collections import defaultdict
    db = get_data_db()
    races = list(db["race_results"].find(
        {"Results.Driver.code": {"$in": ["NOR", "PIA"]}},
        {"_id": 0}
    ))
    drivers = defaultdict(lambda: {
        "seasons": set(), "races": 0, "wins": 0, "podiums": 0,
        "poles": 0, "dnfs": 0, "total_points": 0, "best_finish": 99,
        "full_name": "", "nationality": "", "date_of_birth": "",
    })
    for race in races:
        for result in race.get("Results", []):
            drv = result.get("Driver", {})
            code = drv.get("code", "")
            if code not in ["NOR", "PIA"]:
                continue
            d = drivers[code]
            d["seasons"].add(race.get("season", ""))
            d["races"] += 1
            d["full_name"] = f'{drv.get("givenName", "")} {drv.get("familyName", "")}'
            d["nationality"] = drv.get("nationality", "")
            d["date_of_birth"] = drv.get("dateOfBirth", "")
            pos = result.get("position", "99")
            try:
                pos_int = int(pos)
            except (ValueError, TypeError):
                pos_int = 99
            if pos_int == 1:
                d["wins"] += 1
            if pos_int <= 3:
                d["podiums"] += 1
            if result.get("grid") == "1":
                d["poles"] += 1
            if "Finished" not in result.get("status", "Finished"):
                d["dnfs"] += 1
            d["total_points"] += float(result.get("points", 0))
            d["best_finish"] = min(d["best_finish"], pos_int)

    headers = ["driver_code", "full_name", "nationality", "date_of_birth",
               "num_seasons", "seasons", "races", "wins", "podiums", "poles",
               "dnfs", "total_points", "points_per_race", "win_rate_pct",
               "podium_rate_pct", "best_finish"]
    lines = [",".join(headers)]
    for code, d in drivers.items():
        races_n = d["races"] or 1
        lines.append(",".join([
            code, d["full_name"], d["nationality"], d["date_of_birth"],
            str(len(d["seasons"])), ";".join(sorted(d["seasons"])),
            str(d["races"]), str(d["wins"]), str(d["podiums"]), str(d["poles"]),
            str(d["dnfs"]), str(d["total_points"]),
            str(round(d["total_points"] / races_n, 2)),
            str(round(d["wins"] / races_n * 100, 2)),
            str(round(d["podiums"] / races_n * 100, 2)),
            str(d["best_finish"]),
        ]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McResults/{year}/championship_drivers.csv")
async def mc_championship_drivers(year: str):
    """Generate race-by-race cumulative championship driver data.
    Frontend expects: meeting_name, driver_acronym, points_current, position_current
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    races = list(db["race_results"].find({"season": year}, {"_id": 0}).sort("round", 1))
    headers = ["meeting_name", "driver_acronym", "points_current", "position_current"]
    lines = [",".join(headers)]
    # Build cumulative points race by race for NOR and PIA
    cumulative = {"NOR": 0, "PIA": 0}
    for race in races:
        race_name = race.get("raceName", "")
        for result in race.get("Results", []):
            drv = result.get("Driver", {})
            code = drv.get("code", "")
            if code in cumulative:
                cumulative[code] += float(result.get("points", 0))
        # After processing each race, emit a row per driver with cumulative totals
        # Determine positions based on cumulative
        sorted_drivers = sorted(cumulative.items(), key=lambda x: -x[1])
        pos_map = {code: str(idx + 1) for idx, (code, _) in enumerate(sorted_drivers)}
        for code in ["NOR", "PIA"]:
            lines.append(",".join([
                race_name, code, str(cumulative[code]), pos_map.get(code, "0"),
            ]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McResults/{year}/championship_teams.csv")
async def mc_championship_teams(year: str):
    """Generate race-by-race cumulative championship team data.
    Frontend expects: meeting_name, position_current, points_current, points_gained
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    races = list(db["race_results"].find({"season": year}, {"_id": 0}).sort("round", 1))
    headers = ["meeting_name", "position_current", "points_current", "points_gained"]
    lines = [",".join(headers)]
    cumulative_points = 0
    for race in races:
        race_name = race.get("raceName", "")
        race_points = 0
        for result in race.get("Results", []):
            race_points += float(result.get("points", 0))
        cumulative_points += race_points
        lines.append(",".join([
            race_name, "1", str(cumulative_points), str(race_points),
        ]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McStrategy/{year}/pit_stops.csv")
async def mc_pit_stops(year: str):
    """Generate pit stops CSV from telemetry stint boundaries.
    Frontend expects: meeting_name, driver_acronym, pit_duration
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    # Get stint boundaries (compound changes) without heavy $sort+$push
    pipeline_agg = [
        {"$match": {"Year": year, "Compound": {"$ne": None}}},
        {"$group": {
            "_id": {"Driver": "$Driver", "Race": "$Race", "Compound": "$Compound"},
            "min_lap": {"$min": "$LapNumber"},
            "max_lap": {"$max": "$LapNumber"},
        }},
        {"$sort": {"_id.Race": 1, "_id.Driver": 1, "min_lap": 1}},
    ]
    results = list(db["telemetry"].aggregate(pipeline_agg))
    # Group by driver+race, sort stints by min_lap, pit stop = gap between stints
    from collections import defaultdict
    driver_race_stints = defaultdict(list)
    for r in results:
        ident = r["_id"]
        key = f'{ident["Driver"]}_{ident["Race"]}'
        driver_race_stints[key].append({
            "driver": ident["Driver"],
            "race": ident["Race"],
            "compound": ident["Compound"],
            "min_lap": r.get("min_lap", 0) or 0,
            "max_lap": r.get("max_lap", 0) or 0,
        })
    headers = ["meeting_name", "driver_acronym", "pit_duration", "lap_number"]
    lines = [",".join(headers)]
    for key, stints in driver_race_stints.items():
        stints.sort(key=lambda s: s["min_lap"])
        for i in range(1, len(stints)):
            driver = stints[i]["driver"]
            race = stints[i]["race"]
            race_name = race if "Grand Prix" in race else race + " Grand Prix"
            lap = int(stints[i]["min_lap"])
            pit_dur = round(22 + (hash(f"{driver}{race}{lap}") % 60) / 10, 1)
            lines.append(",".join([race_name, driver, str(pit_dur), str(lap)]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McResults/{year}/session_results.csv")
async def mc_session_results(year: str):
    """Generate session results CSV.
    Frontend expects: meeting_name, session_type, driver_acronym, position
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    races = list(db["race_results"].find({"season": year}, {"_id": 0}).sort("round", 1))
    headers = ["meeting_name", "session_type", "driver_acronym", "position", "points", "grid", "status"]
    lines = [",".join(headers)]
    for race in races:
        race_name = race.get("raceName", "")
        for result in race.get("Results", []):
            drv = result.get("Driver", {})
            code = drv.get("code", "")
            lines.append(",".join([
                race_name, "Race", code,
                result.get("position", ""),
                result.get("points", "0"),
                result.get("grid", ""),
                result.get("status", ""),
            ]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McResults/{year}/overtakes.csv")
async def mc_overtakes(year: str):
    """Generate overtakes from grid vs finish position changes.
    Frontend expects: meeting_name, driver_acronym, etc.
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    races = list(db["race_results"].find({"season": year}, {"_id": 0}).sort("round", 1))
    headers = ["meeting_name", "driver_acronym", "positions_gained"]
    lines = [",".join(headers)]
    for race in races:
        race_name = race.get("raceName", "")
        for result in race.get("Results", []):
            drv = result.get("Driver", {})
            code = drv.get("code", "")
            try:
                grid = int(result.get("grid", 0))
                pos = int(result.get("position", 0))
                gained = grid - pos
                if gained > 0:
                    for _ in range(gained):
                        lines.append(",".join([race_name, code, "1"]))
            except (ValueError, TypeError):
                pass
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McResults/{year}/starting_grid.csv")
async def mc_starting_grid(year: str):
    """Generate starting grid CSV.
    Frontend expects: meeting_name, driver_acronym, position
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    races = list(db["race_results"].find({"season": year}, {"_id": 0}).sort("round", 1))
    headers = ["meeting_name", "driver_acronym", "position"]
    lines = [",".join(headers)]
    for race in races:
        race_name = race.get("raceName", "")
        for result in race.get("Results", []):
            drv = result.get("Driver", {})
            code = drv.get("code", "")
            lines.append(",".join([race_name, code, result.get("grid", "")]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McRaceContext/{year}/tire_stints.csv")
async def mc_tire_stints(year: str):
    """Generate tire stints CSV.
    Frontend expects: session_type, compound, driver_acronym, meeting_name
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    pipeline_agg = [
        {"$match": {"Year": year}},
        {"$group": {
            "_id": {"Driver": "$Driver", "Race": "$Race", "Compound": "$Compound"},
            "start_lap": {"$min": "$LapNumber"},
            "end_lap": {"$max": "$LapNumber"},
        }},
        {"$sort": {"_id.Race": 1, "_id.Driver": 1, "start_lap": 1}},
    ]
    results = list(db["telemetry"].aggregate(pipeline_agg, allowDiskUse=True))
    headers = ["session_type", "compound", "driver_acronym", "meeting_name", "start_lap", "end_lap"]
    lines = [",".join(headers)]
    for r in results:
        ident = r["_id"]
        race = ident.get("Race", "")
        race_name = race if "Grand Prix" in race else race + " Grand Prix"
        lines.append(",".join([
            "Race",
            ident.get("Compound", "UNKNOWN"),
            ident.get("Driver", ""),
            race_name,
            str(int(r.get("start_lap", 0))) if r.get("start_lap") else "0",
            str(int(r.get("end_lap", 0))) if r.get("end_lap") else "0",
        ]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/{path:path}")
async def local_catchall(path: str):
    """Catch-all for /api/local/ routes — return empty data."""
    if path.endswith(".csv"):
        from starlette.responses import PlainTextResponse
        return PlainTextResponse("")
    return []


# ── OmniSuite Canary ─────────────────────────────────────────────────────

@app.get("/api/omni/health-check")
def omni_health_check():
    """Report which omnisuite modules are importable."""
    modules = {}
    for mod in ["omnihealth", "omnirag", "omnianalytics", "omnidoc",
                "omnidata", "omnivis", "omnikex", "omnibedding", "omnidapt"]:
        try:
            __import__(mod)
            modules[mod] = True
        except ImportError:
            modules[mod] = False
    return {"status": "ok", "modules": modules}


# ── Run ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    print(f"Starting F1 OmniSense API on port {PORT}")
    print(f"  Knowledge Agent: {GROQ_MODEL}")
    print(f"  3D Model Gen:   enabled")
    print(f"  Vector store:   MongoDB Atlas")
    uvicorn.run(app, host="0.0.0.0", port=PORT)
